"""
Syntesizer.create_word_doc implementation.

Takes the AlignmentOutput produced by AlignmentOutput.from_result(Alignment.align(...))
and renders a polished Word document: a cover page (with an optional banner /
background image), an auto-generated contents list, and one styled section per
slide — synthesized explanation (Arabic prose that keeps technical/academic
terms in English), the slide's visuals with captions, and the original
OCR/ASR source text tucked into a clearly-labelled "appendix" box so the main
flow stays readable.
"""

import os
import re
import json
import unicodedata
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from wajeezai_api.services.alignment import AlignmentOutput


# ============================================================
# Bidi / RTL helpers (unchanged from the existing codebase)
# ============================================================

RLM = '\u200F'
LRM = '\u200E'


def _is_arabic_char(ch: str) -> bool:
    try:
        return 'ARABIC' in unicodedata.name(ch)
    except ValueError:
        return False


def _split_runs(text: str):
    if not text:
        return []
    runs = []
    current = text[0]
    current_is_ar = _is_arabic_char(text[0])
    for ch in text[1:]:
        is_ar = _is_arabic_char(ch)
        if is_ar == current_is_ar:
            current += ch
        else:
            runs.append((current_is_ar, current))
            current = ch
            current_is_ar = is_ar
    runs.append((current_is_ar, current))
    return runs


def inject_bidi_marks(text: str) -> str:
    """Wraps Latin/number runs inside Arabic text with LRM/RLM so mixed
    English technical terms don't scramble the RTL line order."""
    result = []
    for is_ar, segment in _split_runs(text):
        if is_ar:
            result.append(segment)
        else:
            stripped = segment.strip()
            lead = segment[: len(segment) - len(segment.lstrip())]
            trail = segment[len(segment.rstrip()):]
            result.append(f"{lead}{LRM}{stripped}{RLM}{trail}")
    return ''.join(result)


def add_rtl_paragraph(doc_or_container, text, font_name='Arial', font_size=Pt(12),
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=False,
                       color: RGBColor | None = None):
    fixed_text = inject_bidi_marks(text)
    p = doc_or_container.add_paragraph()
    run = p.add_run(fixed_text)
    p.alignment = alignment
    run.font.rtl = True
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Make sure the complex-script font (used for Arabic) is also set —
    # python-docx's run.font.name only sets the ascii/hAnsi fonts.
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font_name)
    return p


# ============================================================
# Low-level styling helpers (borders, shading, fields, floating images)
# ============================================================

ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # deep academic blue
ACCENT_HEX = "1F4E79"
SOFT_BG_HEX = "EDF2F7"                    # light blue-gray box fill
MUTED_GRAY = RGBColor(0x59, 0x59, 0x59)


def _shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def _box_paragraph(paragraph, color_hex=ACCENT_HEX, size=8, space=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), str(space))
        el.set(qn('w:color'), color_hex)
        pBdr.append(el)
    pPr.append(pBdr)


def _add_page_number_field(paragraph):
    """Inserts a Word PAGE field so the footer stays correct even as content
    reflows (Word computes the number when the document is opened)."""
    run = paragraph.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _float_picture_behind_text(run, page_width_emu, page_height_emu, transparency_pct=0):
    """Converts the inline picture just added via run.add_picture(...) into a
    floating picture anchored to the page, positioned at (0,0) and sized to
    the full page, sitting behind the text — i.e. a page background image.
    Optionally fades it via alphaModFix so body text stays readable.
    """
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))

    extent = inline.find(qn('wp:extent'))
    doc_pr = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))
    cnv_frame_pr = inline.find(qn('wp:cNvGraphicFramePr'))

    # Force the extent to exactly cover the page.
    extent.set('cx', str(page_width_emu))
    extent.set('cy', str(page_height_emu))

    anchor = OxmlElement('wp:anchor')
    anchor.set('behindDoc', '1')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')
    anchor.set('relativeHeight', '1')

    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')
    pos_h_off = OxmlElement('wp:posOffset')
    pos_h_off.text = '0'
    pos_h.append(pos_h_off)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')
    pos_v_off = OxmlElement('wp:posOffset')
    pos_v_off.text = '0'
    pos_v.append(pos_v_off)

    effect_extent = OxmlElement('wp:effectExtent')
    for attr in ('l', 't', 'r', 'b'):
        effect_extent.set(attr, '0')

    wrap_none = OxmlElement('wp:wrapNone')

    anchor.append(simple_pos)
    anchor.append(pos_h)
    anchor.append(pos_v)
    anchor.append(extent)
    anchor.append(effect_extent)
    anchor.append(wrap_none)
    anchor.append(doc_pr)
    if cnv_frame_pr is not None:
        anchor.append(cnv_frame_pr)
    anchor.append(graphic)

    if transparency_pct:
        blip = graphic.find('.//' + qn('a:blip'))
        if blip is not None:
            alpha = OxmlElement('a:alphaModFix')
            alpha.set('amt', str(int((100 - transparency_pct) * 1000)))
            blip.append(alpha)

    drawing.remove(inline)
    drawing.append(anchor)


def add_page_background_image(document, image_path, transparency_pct=70):
    """Adds `image_path` as a faded full-page background on every page, by
    placing it as a behind-text floating picture in the default header
    (headers repeat automatically on every page in the section)."""
    if not image_path or not os.path.exists(image_path):
        return

    section = document.sections[0]
    section.header.is_linked_to_previous = False
    page_w = section.page_width
    page_h = section.page_height

    paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Emu(page_w), height=Emu(page_h))
    _float_picture_behind_text(run, page_w, page_h, transparency_pct=transparency_pct)


def add_cover_banner(document, image_path, max_width_in=6.5):
    """Adds the image inline, full-width-ish, at the top of the cover page
    (a simple, reliable alternative to a true bleed background for the
    title page)."""
    if not image_path or not os.path.exists(image_path):
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(max_width_in))


# ============================================================
# LLM-based per-slide synthesis (Gemini 3.5 flash-lite)
# ============================================================

DEFAULT_SYNTHESIS_MODEL = "gemini-3.5-flash-lite"

_SYNTHESIS_SYSTEM_PROMPT = """أنت مساعد أكاديمي متخصص في تلخيص شرائح محاضرات جامعية (خصوصاً في مجالات علمية/تقنية) وصياغتها كنص عربي فصيح ومترابط.

سيتم تزويدك بـ:
- نصوص مستخرجة من الشريحة (OCR).
- مقاطع من كلام المحاضر المطابقة زمنياً/دلالياً لتلك النصوص (ASR)، وهي غالباً محكية بالعامية وقد تحوي أخطاء تفريغ صوتي.
- أوصاف للرسومات التوضيحية الموجودة في الشريحة (إن وجدت).

قواعد الصياغة (مهمة جداً):
1. اكتب الشرح بالعربية الفصحى الأكاديمية.
2. أي مصطلح علمي أو تقني أو أكاديمي متعارف عليه بالإنجليزية (مثل: Ring Topology, Mesh Topology, Bandwidth, Token Ring, Single Point of Failure, Data Collision) يُكتب كما هو بالإنجليزية داخل الجملة العربية، ولا يُترجم ولا يُعرَّب. مثال صحيح:
   "تتميز طوبولوجيا الـ Ring Topology بأن البيانات تنتقل باتجاه واحد فقط، وأي عطل في الكابل يؤدي إلى توقف الشبكة بالكامل (Single Point of Failure)."
3. لا تُترجم أسماء علمية أو اختصارات (مثل OCR, ASR, LAN) — أبقها كما هي.
4. نقّح كلام المحاضر العامي وحوّله لصياغة أكاديمية واضحة، لكن حافظ على المعنى والمعلومات كما وردت، ولا تخترع معلومات غير موجودة في المصدر.
5. إن وردت معادلة أو قانون رياضي (مثل L = n(n-1)/2) اذكره كما هو بدون تغيير.
6. الأسلوب: فقرة متماسكة (لا نقاط عشوائية) بحد أقصى 5-6 جمل، تشرح المفهوم كأنها جزء من كتاب محاضرات منظم.

أخرج JSON فقط (بدون Markdown ولا أي نص خارج الأقواس) بالشكل التالي:
{
  "title": "عنوان أكاديمي قصير للشريحة (يمكن أن يحوي المصطلح الإنجليزي، مثال: طوبولوجيا الحلقة (Ring Topology))",
  "explanation": "الفقرة الشارحة كما هو موصوف أعلاه"
}
"""


def _build_slide_prompt(slide_output) -> str:
    parts = [f"رقم الشريحة: {slide_output.slide_number}"]

    for i, text_item in enumerate(slide_output.texts, start=1):
        parts.append(f"\n--- نص الشريحة {i} (OCR) ---\n{text_item.ocr_chunk}")
        if text_item.asr_chunks:
            parts.append(f"--- كلام المحاضر المرتبط ---\n{' '.join(text_item.asr_chunks)}")

    for i, visual in enumerate(slide_output.visuals, start=1):
        parts.append(f"\n--- وصف الرسم التوضيحي {i} ---\n{visual.ocr_chunk}")
        if visual.asr_chunks:
            parts.append(f"--- كلام المحاضر المرتبط بالرسم ---\n{' '.join(visual.asr_chunks)}")

    return "\n".join(parts)


def _strip_json_fences(text: str) -> str:
    return re.sub(r"```json|```", "", text).strip()


def synthesize_slide(client, slide_output, model_name: str = DEFAULT_SYNTHESIS_MODEL) -> dict:
    """Calls the LLM to produce {"title": ..., "explanation": ...} for one
    slide. Falls back to a heuristic summary if the client/call fails."""
    prompt = _build_slide_prompt(slide_output)

    try:
        response = client.models.generate_content(
            model=model_name,
            contents=[
                {"role": "user", "parts": [{"text": _SYNTHESIS_SYSTEM_PROMPT}]},
                {"role": "user", "parts": [{"text": prompt}]},
            ],
        )
        cleaned = _strip_json_fences(response.text)
        data = json.loads(cleaned)
        title = data.get("title") or f"الشريحة {slide_output.slide_number}"
        explanation = data.get("explanation") or ""
        print('ronaldo')
        print(title,explanation)
        return {"title": title, "explanation": explanation}
    except Exception as exc:
        print(f"⚠️ Synthesis failed for slide {slide_output.slide_number}: {exc}")
        fallback_bits = [t.ocr_chunk for t in slide_output.texts if t.ocr_chunk]
        fallback_text = " ".join(fallback_bits)[:400]
        return {
            "title": f"الشريحة {slide_output.slide_number}",
            "explanation": fallback_text or "لا يوجد شرح متاح لهذه الشريحة.",
        }


# ============================================================
# Document rendering
# ============================================================

def _add_footer_page_numbers(document):
    section = document.sections[0]
    footer_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = ""
    lead = footer_p.add_run("صفحة ")
    lead.font.size = Pt(9)
    lead.font.color.rgb = MUTED_GRAY
    _add_page_number_field(footer_p)


def _render_cover_page(doc, banner_image_path):
    if banner_image_path:
        add_cover_banner(doc, banner_image_path)
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(inject_bidi_marks('ملخص المحاضرة'))
    title_run.font.rtl = True
    title_run.font.size = Pt(32)
    title_run.bold = True
    title_run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Lecture Notes')
    subtitle_run.font.size = Pt(16)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = MUTED_GRAY

    doc.add_paragraph()
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _box_paragraph(rule, color_hex=ACCENT_HEX, size=12, space=2)
    rule.paragraph_format.space_before = Pt(0)

    doc.add_page_break()


def _render_contents(doc, titles):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_run = heading.add_run(inject_bidi_marks('المحتويات'))
    heading_run.font.rtl = True
    heading_run.bold = True
    heading_run.font.size = Pt(20)
    heading_run.font.color.rgb = ACCENT
    heading.paragraph_format.space_after = Pt(12)

    for idx, title_text in enumerate(titles, start=1):
        line = add_rtl_paragraph(doc, f"{idx}.  {title_text}", font_size=Pt(12))
        line.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def _render_slide_section(doc, slide_output, synthesis, used_image_paths):
    # --- Section heading, numbered like a real lecture-notes chapter ---
    heading = doc.add_heading('', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading_run = heading.add_run(
        inject_bidi_marks(f"{slide_output.slide_number}. {synthesis['title']}")
    )
    heading_run.font.rtl = True
    heading_run.font.color.rgb = ACCENT
    for r in heading.runs:
        r.font.rtl = True

    # --- Synthesized explanation, in a soft shaded box ---
    if synthesis.get("explanation"):
        exp_p = add_rtl_paragraph(doc, synthesis["explanation"], font_size=Pt(12))
        exp_p.paragraph_format.space_before = Pt(4)
        exp_p.paragraph_format.space_after = Pt(10)
        exp_p.paragraph_format.left_indent = Inches(0.15)
        exp_p.paragraph_format.right_indent = Inches(0.15)
        _shade_paragraph(exp_p, SOFT_BG_HEX)
        _box_paragraph(exp_p, color_hex="B8C7DA", size=4, space=6)

    # --- Visuals: image + caption, centered, deduped across the doc ---
    for visual in slide_output.visuals:
        path = visual.path
        caption = (visual.ocr_chunk or '').strip()

        if path and path in used_image_paths:
            continue

        if path and os.path.exists(path):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_before = Pt(6)
            img_run = img_p.add_run()
            img_run.add_picture(path, width=Inches(4.5))
            used_image_paths.add(path)
        elif path:
            print(f"⚠️ Missing image file, skipping: {path}")

        if caption:
            cap_p = add_rtl_paragraph(
                doc, f"شكل توضيحي: {caption}", font_size=Pt(9.5),
                alignment=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED_GRAY,
            )
            cap_p.runs[0].italic = True
            cap_p.paragraph_format.space_after = Pt(10)

class Syntesizer:
    @staticmethod
    def create_word_doc(
        alignment_output: AlignmentOutput,
        client=None,
        output_docx: str = "Lecture_Notes.docx",
        model_name: str = DEFAULT_SYNTHESIS_MODEL,
        cover_image_path: str | None = None,
        background_image_path: str | None = None,
        background_transparency_pct: int = 78,
    ) -> str:
        """
        Creates a Word document from an AlignmentOutput.

        cover_image_path: optional banner image shown on the title page.
        background_image_path: optional image applied as a faded full-page
            background (behind the text) on every page.
        Returns the path to the created Word document.
        """
        slides = alignment_output.slides
        if not slides:
            raise ValueError("alignment_output.slides is empty — nothing to render.")

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Page setup: comfortable margins for an RTL document.
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

        if background_image_path:
            add_page_background_image(
                doc, background_image_path, transparency_pct=background_transparency_pct
            )

        _add_footer_page_numbers(doc)

        _render_cover_page(doc, cover_image_path)

        # --- Pass 1: synthesize every slide up front so we can build a
        #     real contents list before rendering the body. ---
        syntheses = []
        for idx, slide_output in enumerate(slides):
            print(f"Synthesizing slide {idx + 1}/{len(slides)} "
                  f"(slide_number={slide_output.slide_number})...")
            if client is not None:
                synthesis = synthesize_slide(client, slide_output, model_name=model_name)
            else:
                synthesis = {
                    "title": f"الشريحة {slide_output.slide_number}",
                    "explanation": " ".join(
                        t.ocr_chunk for t in slide_output.texts if t.ocr_chunk
                    )[:400],
                }
            syntheses.append(synthesis)

        _render_contents(doc, [s["title"] for s in syntheses])

        # --- Pass 2: render each slide's section. ---
        used_image_paths = set()
        for idx, (slide_output, synthesis) in enumerate(zip(slides, syntheses)):
            _render_slide_section(doc, slide_output, synthesis, used_image_paths)
            if idx < len(slides) - 1:
                doc.add_page_break()

        doc.save(output_docx)
        print(f"Success! Document saved to {output_docx}")
        return output_docx

    @staticmethod
    def create_pdf(
        alignment_output: AlignmentOutput,
        client=None,
        output_pdf: str = "Lecture_Notes.pdf",
        model_name: str = DEFAULT_SYNTHESIS_MODEL,
        cover_image_path: str | None = None,
        background_image_path: str | None = None,
        background_transparency_pct: int = 78,
        libreoffice_path: str = "soffice",
    ) -> str:
        """
        Creates a PDF from AlignmentOutput.

        The document is first rendered as DOCX using create_word_doc(),
        then converted to PDF using LibreOffice.

        Returns the path to the created PDF.
        """

        output_pdf = Path(output_pdf)
        output_pdf.parent.mkdir(parents=True, exist_ok=True)

        # Temporary DOCX next to the target PDF
        docx_path = output_pdf.with_suffix(".docx")

        # Create the Word document
        Syntesizer.create_word_doc(
            alignment_output=alignment_output,
            client=client,
            output_docx=str(docx_path),
            model_name=model_name,
            cover_image_path=cover_image_path,
            background_image_path=background_image_path,
            background_transparency_pct=background_transparency_pct,
        )

        # Convert DOCX -> PDF
        try:
            subprocess.run(
                [
                    libreoffice_path,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_pdf.parent),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )

        except FileNotFoundError:
            raise RuntimeError(
                "LibreOffice was not found. "
                "Install LibreOffice or provide the correct "
                "path through libreoffice_path."
            )

        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"LibreOffice PDF conversion failed:\n"
                f"{e.stdout}\n"
                f"{e.stderr}"
            )

        generated_pdf = output_pdf.parent / f"{docx_path.stem}.pdf"

        if not generated_pdf.exists():
            raise RuntimeError(
                f"PDF conversion completed but the PDF was not found: "
                f"{generated_pdf}"
            )

        print(f"Success! PDF saved to {generated_pdf}")

        return str(generated_pdf)